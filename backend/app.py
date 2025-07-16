from utils import aiprocess, generate_brief_from_file, edit_pdf, search_google
import os
from flask import Flask, request, jsonify, send_file
import pymysql
import io
import os
from docx import Document
import fitz  # PyMuPDF
import PyPDF2
import requests
from bs4 import BeautifulSoup
from serpapi import GoogleSearch
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from flask_cors import CORS
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import textwrap

app = Flask(__name__)
CORS(app)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
db_config = {
    "host": "localhost",
    "user": "root",
    "password": "aldina123",
    "database": "cat"
}

GOOGLE_API_KEY = "AIzaSyBDidg9orPvjjQfMz6n1tNx8RWgLjEipeQ"
SERP_API_KEY = "dad8d20aacff98df37793a921fe61fad4ddadfc0d2714150c739529c8b0e3c2c"

def create_pdf_from_text(text, filename):
    """Create a PDF from text content"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Split text into paragraphs
    paragraphs = text.split('\n')
    
    for para in paragraphs:
        if para.strip():  # Skip empty lines
            # Wrap long lines
            wrapped_lines = textwrap.fill(para, width=80)
            p = Paragraph(wrapped_lines, styles['Normal'])
            story.append(p)
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_from_text(text, filename):
    """Create a DOCX from text content"""
    buffer = io.BytesIO()
    doc = Document()
    
    # Split text into paragraphs
    paragraphs = text.split('\n')
    
    for para in paragraphs:
        if para.strip():  # Skip empty lines
            doc.add_paragraph(para)
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_txt_from_text(text, filename):
    """Create a TXT from text content"""
    buffer = io.BytesIO()
    buffer.write(text.encode('utf-8'))
    buffer.seek(0)
    return buffer

@app.route('/upload', methods=['POST'])
def upload_file():
    files = request.files.getlist('files')
    if not files:
        return jsonify({"error": "No files uploaded"}), 400

    paths = []
    uploaded_db_info = []
    conn = None
    cursor = None

    try:
        conn = pymysql.connect(**db_config)
        cursor = conn.cursor()

        for file in files:
            if file and file.filename:
                filename = file.filename

                # 1. Save to disk
                path = os.path.join(UPLOAD_FOLDER, filename)
                file.seek(0)
                file.save(path)
                paths.append(path)

                # 2. Save to database
                file.seek(0)
                content = file.read()
                insert_query = "INSERT INTO files (filename, content) VALUES (%s, %s)"
                cursor.execute(insert_query, (filename, content))
                conn.commit()
                file_id = cursor.lastrowid

                uploaded_db_info.append({
                    "filename": filename,
                    "file_id": file_id
                })

        return jsonify({
            "message": f"{len(paths)} files uploaded",
            "paths": paths,
            "db_records": uploaded_db_info
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@app.route('/brief', methods=['POST'])
def brief():
    file_path = request.json.get('file_path')
    if not file_path:
        return jsonify({"error": "file_path required"}), 400
    summary = generate_brief_from_file(file_path)
    return jsonify({"summary": summary})

@app.route('/ask', methods=['POST'])
def ask():
    data = request.json
    file_paths = data.get("file_paths", [])
    question = data.get("question", "")
    if not file_paths or not question:
        return jsonify({"error": "Missing file_paths or question"}), 400
    answer = aiprocess(file_paths, question)
    return jsonify({"answer": answer})

@app.route('/edit-file/<int:file_id>', methods=['POST'])
def edit_file(file_id):
    data = request.json
    filename = data.get("filename")
    contents = data.get("contents")

    if not filename:
        return jsonify({"error": "Missing filename"}), 400
    if not contents:
        return jsonify({"error": "Missing contents"}), 400

    # Initialize variables for cleanup
    conn = None
    cursor = None
    temp_path = None

    try:
        # Get file extension
        ext = os.path.splitext(filename)[1].lower()
        
        # Create the modified file based on extension
        if ext == '.pdf':
            modified_file = create_pdf_from_text(contents, filename)
            content_type = 'application/pdf'
        elif ext == '.docx':
            modified_file = create_docx_from_text(contents, filename)
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif ext == '.txt':
            modified_file = create_txt_from_text(contents, filename)
            content_type = 'text/plain'
        else:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 400
        
        # Update database with modified file
        conn = pymysql.connect(**db_config)
        cursor = conn.cursor()
        modified_file.seek(0)
        content = modified_file.read()
        update_query = "UPDATE files SET filename = %s, content = %s WHERE id = %s"
        new_filename = f"edited_{filename}"
        cursor.execute(update_query, (new_filename, content, file_id))
        conn.commit()

        if cursor.rowcount == 0:
            return jsonify({"error": "No file found with given file_id"}), 404

        # Return success response with file info
        return jsonify({
            "message": "File edited successfully",
            "file_id": file_id,
            "original_filename": filename,
            "new_filename": new_filename,
            "file_type": ext,
            "content_type": content_type
        })

    except Exception as e:
        return jsonify({"error": "Failed to edit file", "details": str(e)}), 500
    finally:
        # Clean up resources
        if cursor:
            cursor.close()
        if conn:
            conn.close()
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)

@app.route('/edit-file-download/<int:file_id>', methods=['POST'])
def edit_file_download(file_id):
    data = request.json
    filename = data.get("filename")
    contents = data.get("contents")

    if not filename:
        return jsonify({"error": "Missing filename"}), 400
    if not contents:
        return jsonify({"error": "Missing contents"}), 400

    # Initialize variables for cleanup
    conn = None
    cursor = None
    temp_path = None

    try:
        # Get file extension
        ext = os.path.splitext(filename)[1].lower()
        
        # Create the modified file based on extension
        if ext == '.pdf':
            modified_file = create_pdf_from_text(contents, filename)
            mimetype = 'application/pdf'
            download_name = f"edited_{filename}"
        elif ext == '.docx':
            modified_file = create_docx_from_text(contents, filename)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            download_name = f"edited_{filename}"
        elif ext == '.txt':
            modified_file = create_txt_from_text(contents, filename)
            mimetype = 'text/plain'
            download_name = f"edited_{filename}"
        else:
            return jsonify({"error": f"Unsupported file type: {ext}"}), 400
        
        # Update database with modified file
        conn = pymysql.connect(**db_config)
        cursor = conn.cursor()
        modified_file.seek(0)
        content = modified_file.read()
        update_query = "UPDATE files SET filename = %s, content = %s WHERE id = %s"
        new_filename = f"edited_{filename}"
        cursor.execute(update_query, (new_filename, content, file_id))
        conn.commit()

        if cursor.rowcount == 0:
            return jsonify({"error": "No file found with given file_id"}), 404

        # Return the modified file for download
        modified_file.seek(0)
        return send_file(
            modified_file, 
            download_name=download_name, 
            as_attachment=True,
            mimetype=mimetype
        )

    except Exception as e:
        return jsonify({"error": "Failed to edit file", "details": str(e)}), 500
    finally:
        # Clean up resources
        if cursor:
            cursor.close()
        if conn:
            conn.close()
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)

@app.route('/get-file-content/<int:file_id>', methods=['GET'])
def get_file_content(file_id):
    conn = None
    cursor = None

    try:
        conn = pymysql.connect(**db_config)
        cursor = conn.cursor()
        cursor.execute("SELECT filename, content FROM files WHERE id = %s", (file_id,))
        result = cursor.fetchone()

        if not result:
            return jsonify({"error": "File not found"}), 404

        filename, content = result
        ext = os.path.splitext(filename)[1].lower()
        temp_path = f"temp_{file_id}{ext}"

        # Write content to temporary file
        with open(temp_path, "wb") as f:
            f.write(content)

        # Extract text based on file type
        if ext == ".pdf":
            doc = fitz.open(temp_path)
            text = "".join(page.get_text() for page in doc)
            doc.close()

        elif ext == ".docx":
            doc = Document(temp_path)
            text = "\n".join(para.text for para in doc.paragraphs)

        elif ext == ".txt":
            with open(temp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()

        else:
            os.remove(temp_path)
            return jsonify({"filename": filename, "message": f"Unsupported file type: {ext}"}), 415

        os.remove(temp_path)

        return jsonify({
            "file_id": file_id,
            "filename": filename,
            "content": text
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@app.route('/websearch', methods=['POST'])
def web_search():
    data = request.json
    query = data.get("query")

    if not query:
        return jsonify({"error": "Query is required"}), 400

    try:
        # Step 1: Get search result links using utils function
        links = search_google(query)

        # Step 2: Scrape content from each link
        snippets = []
        headers = {"User-Agent": "Mozilla/5.0"}

        for url in links[:5]:  # Limit to top 5 for performance
            try:
                response = requests.get(url, headers=headers, timeout=10)
                soup = BeautifulSoup(response.text, "html.parser")

                # Extract visible paragraph texts
                text = " ".join(p.text.strip() for p in soup.find_all("p") if p.text.strip())
                snippet = text[:1000] + "..." if len(text) > 1000 else text

                snippets.append({
                    "url": url,
                    "snippet": snippet
                })
            except Exception as e:
                snippets.append({
                    "url": url,
                    "snippet": f"Failed to fetch content: {str(e)}"
                })

        return jsonify({
            "query": query,
            "results": snippets
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/files', methods=['GET'])
def get_all_files():
    conn = None
    cursor = None

    try:
        conn = pymysql.connect(**db_config)
        cursor = conn.cursor()
        cursor.execute("SELECT id, filename FROM files")
        results = cursor.fetchall()

        files = [{"id": row[0], "filename": row[1]} for row in results]

        return jsonify({
            "count": len(files),
            "files": files
        }), 200

    except Exception as e:
        print("Error in /files:", e)
        return jsonify({"error": str(e)}), 500

    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

if __name__ == '__main__':
    app.run(debug=True)