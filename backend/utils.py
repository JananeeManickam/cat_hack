import os
import io
import json
import requests
from bs4 import BeautifulSoup
from docx import Document
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.prompts import PromptTemplate
from langchain_community.vectorstores import Chroma
# from serpapi import GoogleSearch
import google.generativeai as genai
import urllib.parse
import pandas as pd

# ---------- PDF Editing ----------
def edit_pdf(file_storage, text_to_add="Modified"):
    """
    Edits a PDF by adding metadata. Accepts a Flask `request.files["file"]` object.
    """
    if file_storage is None:
        raise ValueError("No file provided")

    # Read the file content into a BytesIO buffer
    file_stream = io.BytesIO(file_storage.read())
    file_stream.seek(0)

    # Use PyPDF2 to read and modify the PDF
    reader = PdfReader(file_stream)
    writer = PdfWriter()

    for page in reader.pages:
        writer.add_page(page)

    writer.add_metadata({"/ModifiedBy": text_to_add})

    output_stream = io.BytesIO()
    writer.write(output_stream)
    output_stream.seek(0)

    return output_stream

# ---------- Brief Summary Generator ----------
def generate_brief_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return "This is a PDF document. Contents: (simulated brief summary)."
    elif ext in ['.docx', '.doc']:
        return "This is a Word document. Contents: (simulated brief summary)."
    elif ext == '.xlsx':
        return "This is an Excel file. Contents: (simulated brief summary)."
    elif ext == '.csv':
        return "This is a CSV file. Contents: (simulated brief summary)."
    else:
        return "Unsupported file format for summarization."

# ---------- File Text Extraction ----------
def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as f:
        pdf_reader = PyPDF2.PdfReader(f)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return "\n\n".join(para.text for para in doc.paragraphs)

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding='utf-8') as f:
        return f.read()

# ---------- AI Processing with Gemini ----------
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.prompts import PromptTemplate

def extract_text_from_pdf(path):
    from PyPDF2 import PdfReader
    reader = PdfReader(path)
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def extract_text_from_docx(path):
    from docx import Document
    doc = Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_txt(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def extract_text_from_csv(path):
    print(path,"path")
    df = pd.read_csv(path)
    return f"CSV FILE: {path}\n" + df.to_string(index=False)

def aiprocess(file_paths, question):
    texts = []
    
    for path in file_paths:
        if path.endswith(".pdf"):
            texts.append(extract_text_from_pdf(path))
        elif path.endswith(".docx"):
            texts.append(extract_text_from_docx(path))
        elif path.endswith(".txt"):
            texts.append(extract_text_from_txt(path))
        elif path.endswith(".csv"):
            texts.append(extract_text_from_csv(path))
        else:
            continue  # Unsupported

    context = "\n\n".join(texts)

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(context)

    model = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        google_api_key="AIzaSyBDidg9orPvjjQfMz6n1tNx8RWgLjEipeQ",
        temperature=0.1,
        convert_system_message_to_human=True
    )

    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key="AIzaSyBDidg9orPvjjQfMz6n1tNx8RWgLjEipeQ"
    )

    vector_index = Chroma.from_texts(chunks, embeddings).as_retriever(search_kwargs={"k": 5})

    prompt_template = """You are a highly experienced machine operator with 25 years of expertise.
You have access to terrain data, climate readings, expected machine behavior, and current sensor readings.
Analyze the data provided and determine:

1. Whether any machine is likely to fail or already failing.
2. What are the probable causes?
3. What are the step-by-step recommendations or actions the operator should take to prevent or respond to failure?
4. Make your suggestions practical and safety-first.

only give the final improvements , what isthe cause of failure if any do not give the questions in the response

Context:
{context}

Question: {question}

Helpful Answer (like an expert operator):"""

    QA_CHAIN_PROMPT = PromptTemplate.from_template(prompt_template)

    qa_chain = RetrievalQA.from_chain_type(
        llm=model,
        retriever=vector_index,
        return_source_documents=True,
        chain_type_kwargs={"prompt": QA_CHAIN_PROMPT}
    )

    result = qa_chain({"query": question})
    return  result["result"] 
    

# ---------- Google Search ----------
# def search_google(query):
#     genai.configure(api_key="AIzaSyBj7ruubVa72IK-7RcGFJjHuyerBlveDqI")

#     params = {
#         "q": query,
#         "api_key": "dad8d20aacff98df37793a921fe61fad4ddadfc0d2714150c739529c8b0e3c2c",
#         "num": 15
#     }

#     search = GoogleSearch(params)
#     results = search.get_dict()
#     return [result['link'] for result in results.get("organic_results", [])]
def search_google(query):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        encoded_query = urllib.parse.quote_plus(query)
        url = f"https://www.google.com/search?q={encoded_query}"

        response = requests.get(url, headers=headers)
        soup = BeautifulSoup(response.text, "html.parser")

        links = []
        for a in soup.select('a'):
            href = a.get('href')
            if href and "/url?q=" in href:
                link = href.split("/url?q=")[1].split("&")[0]
                if "google.com" not in link:  # skip internal links
                    links.append(link)

        return links[:10]  # return top 10 results

    except Exception as e:
        print(f"Error searching Google: {e}")
        return []
# ---------- Web Text Extraction ----------
def extract_text_from_url(url):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(url, headers=headers, timeout=10)
        soup = BeautifulSoup(response.text, "html.parser")
        return " ".join([p.text for p in soup.find_all("p")])
    except:
        return ""

